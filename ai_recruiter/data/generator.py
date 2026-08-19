"""Генератор демо-базы резюме.

Создаёт N синтетических резюме (PDF + DOCX) с фиксированным seed, а также
манифест JSON с ground-truth данными (навыки, опыт, seniority) для оценки
качества скоринга.

Запуск:
    python -m ai_recruiter.data.generator --n 100
"""

from __future__ import annotations

import argparse
import json
import random
from dataclasses import asdict, dataclass, field
from pathlib import Path

from ai_recruiter.config import config

# --------------------------------------------------------------------------- #
# Словари данных
# --------------------------------------------------------------------------- #

FIRST_NAMES = [
    "Александр", "Дмитрий", "Иван", "Михаил", "Артём", "Никита", "Егор",
    "Максим", "Павел", "Сергей", "Андрей", "Владимир", "Тимофей", "Кирилл",
    "Анна", "Мария", "Елена", "Ольга", "Наталья", "Дарья", "Екатерина",
    "Виктория", "Алина", "Полина", "Ксения", "Ирина", "Светлана", "Юлия",
]

LAST_NAMES = [
    "Иванов", "Петров", "Смирнов", "Кузнецов", "Соколов", "Попов", "Лебедев",
    "Козлов", "Новиков", "Морозов", "Волков", "Соловьёв", "Васильев", "Зайцев",
    "Иванова", "Петрова", "Смирнова", "Кузнецова", "Соколова", "Попова",
    "Лебедева", "Козлова", "Новикова", "Морозова", "Волкова", "Васильева",
]

PATRONYMICS = [
    "Александрович", "Дмитриевич", "Иванович", "Михайлович", "Сергеевич",
    "Андреевич", "Владимирович", "Николаевич", "Павлович", "Петрович",
    "Александровна", "Дмитриевна", "Ивановна", "Михайловна", "Сергеевна",
    "Андреевна", "Владимировна", "Николаевна", "Павловна", "Петровна",
]

COMPANIES = [
    "Яндекс", "VK", "Сбер", "Тинькофф", "Ozon", "Wildberries", "Авито",
    "Kaspersky", "Positive Technologies", "X5 Tech", "MTS Digital", "Билайн",
    "СберЗдоровье", "Lamoda", "2ГИС", "HeadHunter", "Модульбанк", "Точка",
    "Selectel", "Veeam", "Циан", "Додо Пицца", "ВкусВилл", "Росбанк",
]

UNIVERSITIES = [
    "МГУ им. М.В. Ломоносова", "МФТИ", "ВШЭ", "СПбГУ", "ИТМО", "МГТУ им. Баумана",
    "НИЯУ МИФИ", "УрФУ", "КФУ", "НГУ", "ТГУ", "МИСиС", "МАИ", "Сколтех",
]

ROLES = {
    "Backend-разработчик (Python)": {
        "skills": [
            "Python", "Django", "FastAPI", "PostgreSQL", "Redis", "Docker",
            "Kubernetes", "gRPC", "REST API", "SQLAlchemy", "Celery", "Kafka",
            "asyncio", "pytest", "Git", "Linux", "CI/CD", "MongoDB",
        ],
    },
    "Frontend-разработчик (React)": {
        "skills": [
            "JavaScript", "TypeScript", "React", "Redux", "Next.js", "HTML",
            "CSS", "Sass", "Webpack", "Vite", "Jest", "Storybook", "Node.js",
            "GraphQL", "REST API", "Git", "Docker",
        ],
    },
    "ML-инженер": {
        "skills": [
            "Python", "PyTorch", "TensorFlow", "scikit-learn", "Pandas", "NumPy",
            "NLP", "Transformers", "LLM", "MLflow", "Docker", "Kubernetes",
            "SQL", "Spark", "ONNX", "Computer Vision", "MLOps", "Git",
        ],
    },
    "DevOps-инженер": {
        "skills": [
            "Linux", "Docker", "Kubernetes", "Terraform", "Ansible", "CI/CD",
            "GitLab CI", "GitHub Actions", "AWS", "GCP", "Prometheus", "Grafana",
            "Bash", "Python", "Helm", "Nginx", "PostgreSQL",
        ],
    },
    "Data Analyst": {
        "skills": [
            "SQL", "Python", "Pandas", "NumPy", "Tableau", "Power BI", "Excel",
            "A/B-тесты", "Статистика", "Математическая статистика", "Airflow",
            "ClickHouse", "Redash", "Jupyter", "Визуализация данных",
        ],
    },
    "QA-инженер": {
        "skills": [
            "Python", "pytest", "Selenium", "Playwright", "Postman", "REST API",
            "SQL", "Тестирование API", "Автотесты", "Jenkins", "Docker",
            "Linux", "Jira", "TestRail", "Git", "CI/CD",
        ],
    },
}

SENIORITY_YEARS = {
    "junior": (0.5, 2.0),
    "middle": (2.0, 5.0),
    "senior": (5.0, 12.0),
}

SUMMARY_TEMPLATES = [
    "{role} с опытом {years:.0f} лет. Специализируюсь на {focus}. "
    "Работаю в продуктовых командах, отвечаю за {responsibility}.",
    "Целеустремлённый {role}. {years:.0f} лет коммерческой разработки. "
    "Глубоко погружаюсь в {focus}, стремлюсь к измеримым результатам и качеству.",
    "Опытный специалист в роли {role}. За плечами {years:.0f} лет работы "
    "над {focus}. Умею выстраивать процессы и доносить решения до бизнеса.",
]

EXP_DUTIES = {
    "Backend-разработчик (Python)": [
        "Разрабатывал и поддерживал высоконагруженные REST/gRPC-сервисы.",
        "Проектировал схемы БД и оптимизировал тяжёлые SQL-запросы.",
        "Покрывал код автотестами, настраивал CI/CD и мониторинг.",
        "Интегрировался с внешними API и очередями сообщений.",
    ],
    "Frontend-разработчик (React)": [
        "Разрабатывал SPA на React/TypeScript, внедрял Next.js.",
        "Оптимизировал производительность и переиспользуемые компоненты.",
        "Писал unit-тесты на Jest и интеграционные на Playwright.",
        "Внедрял SSR и кастомизировал сборку на Vite/Webpack.",
    ],
    "ML-инженер": [
        "Обучал и деплоил ML-модели в продакшн.",
        "Строил пайплайны обработки данных и feature engineering.",
        "Тюнил LLM и строил RAG-системы для доменных задач.",
        "Настраивал MLflow и мониторинг качества моделей.",
    ],
    "DevOps-инженер": [
        "Автоматизировал инфраструктуру через Terraform и Ansible.",
        "Строил CI/CD-пайплайны и ускорял доставку релизов.",
        "Настраивал Kubernetes-кластеры и observability-стек.",
        "Снижал стоимость инфраструктуры и повышал её надёжность.",
    ],
    "Data Analyst": [
        "Строил дашборды и отчётность для продуктовых команд.",
        "Проводил A/B-тесты и анализировал их результаты.",
        "Писал SQL-запросы к большим объёмам данных.",
        "Автоматизировал рутинную аналитику на Python.",
    ],
    "QA-инженер": [
        "Разрабатывал автотесты API и UI.",
        "Вёл тестовую документацию и баг-репорты.",
        "Настраивал запуски автотестов в CI.",
        "Проводил нагрузочное и регрессионное тестирование.",
    ],
}

ACHIEVEMENTS = [
    "Сократил время обработки запросов на 40%.",
    "Повысил покрытие автотестами до 85%.",
    "Вывел продукт на рынок в срок, сократив затраты на 20%.",
    "Автоматизировал рутинные процессы, сэкономив ~10 часов в неделю.",
    "Мигрировал легаси-сервис на новую архитектуру без даунтайма.",
    "Настроил мониторинг, снизив число инцидентов вдвое.",
]

LANGUAGES = [
    "Русский — родной, английский — B2",
    "Русский — родной, английский — C1",
    "Русский — родной, английский — B1",
    "Русский — родной, английский — Upper-Intermediate",
    "Русский — родной, английский — Intermediate",
]


# --------------------------------------------------------------------------- #
# Модель данных кандидата
# --------------------------------------------------------------------------- #

@dataclass
class Experience:
    company: str
    position: str
    start: str
    end: str
    duties: list[str] = field(default_factory=list)
    achievements: list[str] = field(default_factory=list)


@dataclass
class Education:
    institution: str
    degree: str
    year: str


@dataclass
class Candidate:
    candidate_id: str
    full_name: str
    role: str
    seniority: str
    years_experience: float
    summary: str
    skills: list[str]
    experience: list[Experience]
    education: list[Education]
    languages: str
    email: str
    resume_file: str = ""


# --------------------------------------------------------------------------- #
# Генерация
# --------------------------------------------------------------------------- #

def _career_timeline(rng: random.Random, years: float) -> list[tuple[str, str]]:
    """Строит непротиворечивый список (start, end) по убыванию от настоящего."""
    now_year, now_month = 2025, 8
    segments = rng.randint(1, 4) if years < 3 else rng.randint(2, 5)
    total_months = int(years * 12)
    remaining = total_months

    timeline: list[tuple[str, str]] = []
    cursor_y, cursor_m = now_year, now_month
    for i in range(segments):
        if remaining <= 0:
            break
        is_last = (i == segments - 1) or (remaining <= 0)
        if is_last:
            dur = remaining
        else:
            dur = max(4, int(remaining * rng.uniform(0.3, 0.6)))
        dur = min(dur, remaining)
        end_y = cursor_y
        end_m = cursor_m
        # у последнего (текущего) места end — «наст. время»
        start_total = (end_y * 12 + end_m) - dur
        start_y, start_m = divmod(start_total, 12)
        if start_m == 0:
            start_m = 12
            start_y -= 1
        end = "наст. время" if i == 0 else f"{end_m:02d}.{end_y}"
        timeline.append((f"{start_m:02d}.{start_y}", end))
        cursor_y, cursor_m = start_y, start_m - 1
        if cursor_m <= 0:
            cursor_m = 12
            cursor_y -= 1
        remaining -= dur

    return timeline


def generate_candidate(rng: random.Random, idx: int) -> Candidate:
    role = rng.choice(list(ROLES.keys()))
    role_skills = ROLES[role]["skills"]
    seniority = rng.choices(
        ["junior", "middle", "senior"], weights=[0.2, 0.5, 0.3], k=1
    )[0]

    lo, hi = SENIORITY_YEARS[seniority]
    years = round(rng.uniform(lo, hi), 1)

    # Имя с согласованием по роду (женская фамилия заканчивается на -а)
    gender = rng.choice(["m", "f"])
    first = rng.choice(FIRST_NAMES)
    if gender == "m":
        last = rng.choice([n for n in LAST_NAMES if not n.endswith("а")])
        patr = rng.choice([p for p in PATRONYMICS if p.endswith("ич")])
    else:
        last = rng.choice([n for n in LAST_NAMES if n.endswith("а")])
        patr = rng.choice([p for p in PATRONYMICS if p.endswith("на")])
    full_name = f"{last} {first} {patr}"

    n_skills = rng.randint(6, 10)
    skills = rng.sample(role_skills, k=min(n_skills, len(role_skills)))

    summary = rng.choice(SUMMARY_TEMPLATES).format(
        role=role,
        years=max(years, 1),
        focus=rng.choice(skills[:4]),
        responsibility=rng.choice(EXP_DUTIES[role]).rstrip(".").lower(),
    )

    timeline = _career_timeline(rng, years)
    experience: list[Experience] = []
    for start, end in timeline:
        n_duties = rng.randint(2, 4)
        n_ach = rng.randint(0, 2)
        experience.append(
            Experience(
                company=rng.choice(COMPANIES),
                position=role if rng.random() < 0.6 else f"{rng.choice(['Старший', 'Ведущий']) if seniority != 'junior' else ''} {role}".strip(),
                start=start,
                end=end,
                duties=rng.sample(EXP_DUTIES[role], k=min(n_duties, len(EXP_DUTIES[role]))),
                achievements=rng.sample(ACHIEVEMENTS, k=n_ach),
            )
        )

    grad_year = 2025 - int(years) - rng.randint(4, 6)
    education = [
        Education(
            institution=rng.choice(UNIVERSITIES),
            degree=rng.choice(["Бакалавр", "Специалист", "Магистр"]),
            year=str(grad_year),
        )
    ]

    email = f"{first.lower()}.{last.lower().replace('ё', 'e')}{rng.randint(1, 999)}@example.com"

    return Candidate(
        candidate_id=f"cand_{idx:04d}",
        full_name=full_name,
        role=role,
        seniority=seniority,
        years_experience=years,
        summary=summary,
        skills=skills,
        experience=experience,
        education=education,
        languages=rng.choice(LANGUAGES),
        email=email,
    )


# --------------------------------------------------------------------------- #
# Рендер в DOCX и PDF
# --------------------------------------------------------------------------- #

def _candidate_to_text(c: Candidate) -> str:
    lines = [c.full_name, c.role, f"Seniority: {c.seniority}", "", c.summary, ""]
    lines.append("Навыки: " + ", ".join(c.skills))
    lines.append("")
    lines.append("Опыт работы:")
    for e in c.experience:
        lines.append(f"- {e.position} — {e.company} ({e.start} — {e.end})")
        for d in e.duties:
            lines.append(f"  * {d}")
        for a in e.achievements:
            lines.append(f"  * Достижение: {a}")
    lines.append("")
    lines.append("Образование:")
    for ed in c.education:
        lines.append(f"- {ed.institution}, {ed.degree}, {ed.year}")
    lines.append("")
    lines.append(f"Языки: {c.languages}")
    lines.append(f"Email: {c.email}")
    return "\n".join(lines)


def _render_docx(c: Candidate, path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(c.full_name, level=1)
    doc.add_paragraph(c.role)
    doc.add_paragraph(f"Seniority: {c.seniority}")
    doc.add_paragraph(c.summary)

    doc.add_heading("Навыки", level=2)
    doc.add_paragraph(", ".join(c.skills))

    doc.add_heading("Опыт работы", level=2)
    for e in c.experience:
        doc.add_paragraph(
            f"{e.position} — {e.company} ({e.start} — {e.end})", style="List Bullet"
        )
        for d in e.duties:
            doc.add_paragraph(d, style="List Bullet 2")
        for a in e.achievements:
            doc.add_paragraph(f"Достижение: {a}", style="List Bullet 2")

    doc.add_heading("Образование", level=2)
    for ed in c.education:
        doc.add_paragraph(f"{ed.institution}, {ed.degree}, {ed.year}", style="List Bullet")

    doc.add_heading("Дополнительно", level=2)
    doc.add_paragraph(f"Языки: {c.languages}")
    doc.add_paragraph(f"Email: {c.email}")

    doc.save(path)


def _find_cyrillic_font() -> str | None:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",  # macOS
        "/Library/Fonts/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",       # Linux
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None


def _render_pdf(c: Candidate, path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_path = _find_cyrillic_font()
    if font_path is None:
        raise RuntimeError("Не найден TTF-шрифт с поддержкой кириллицы.")

    pdfmetrics.registerFont(TTFont("Cyr", font_path))
    normal = ParagraphStyle("normal", fontName="Cyr", fontSize=10, leading=14)
    heading = ParagraphStyle(
        "heading", fontName="Cyr", fontSize=14, leading=18, spaceAfter=6
    )
    sub = ParagraphStyle(
        "sub", fontName="Cyr", fontSize=12, leading=16, spaceBefore=8, spaceAfter=4
    )

    doc = SimpleDocTemplate(str(path), pagesize=A4, topMargin=15 * mm, bottomMargin=15 * mm)
    story = [
        Paragraph(c.full_name, heading),
        Paragraph(c.role, normal),
        Paragraph(f"Seniority: {c.seniority}", normal),
        Spacer(1, 6),
        Paragraph(c.summary, normal),
    ]

    story.append(Paragraph("Навыки", sub))
    story.append(Paragraph(", ".join(c.skills), normal))

    story.append(Paragraph("Опыт работы", sub))
    for e in c.experience:
        story.append(
            Paragraph(f"<b>{e.position} — {e.company}</b> ({e.start} — {e.end})", normal)
        )
        for d in e.duties:
            story.append(Paragraph(f"• {d}", normal))
        for a in e.achievements:
            story.append(Paragraph(f"• Достижение: {a}", normal))

    story.append(Paragraph("Образование", sub))
    for ed in c.education:
        story.append(Paragraph(f"• {ed.institution}, {ed.degree}, {ed.year}", normal))

    story.append(Paragraph("Дополнительно", sub))
    story.append(Paragraph(f"Языки: {c.languages}", normal))
    story.append(Paragraph(f"Email: {c.email}", normal))

    doc.build(story)


# --------------------------------------------------------------------------- #
# Публичный API
# --------------------------------------------------------------------------- #

def _manifest_entry(c: Candidate, filename: str) -> dict:
    return {
        "candidate_id": c.candidate_id,
        "full_name": c.full_name,
        "role": c.role,
        "seniority": c.seniority,
        "years_experience": c.years_experience,
        "skills": c.skills,
        "experience": [asdict(e) for e in c.experience],
        "education": [asdict(ed) for ed in c.education],
        "languages": c.languages,
        "email": c.email,
        "resume_file": filename,
    }


def generate(n: int | None = None, seed: int | None = None) -> list[dict]:
    """Генерирует резюме и манифест, возвращает список записей манифеста."""
    n = n or config.n_resumes
    seed = seed if seed is not None else config.seed
    rng = random.Random(seed)

    config.ensure_dirs()
    manifest = []

    for idx in range(n):
        c = generate_candidate(rng, idx)
        stem = f"{c.candidate_id}_{c.full_name.replace(' ', '_')}"
        pdf_path = config.resumes_dir / f"{stem}.pdf"
        docx_path = config.resumes_dir / f"{stem}.docx"
        _render_pdf(c, pdf_path)
        _render_docx(c, docx_path)
        manifest.append(_manifest_entry(c, f"{stem}.pdf"))

    with open(config.manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    return manifest


def main() -> None:
    parser = argparse.ArgumentParser(description="Генерация демо-базы резюме")
    parser.add_argument("--n", type=int, default=None, help="Количество резюме")
    parser.add_argument("--seed", type=int, default=None, help="Seed для воспроизводимости")
    args = parser.parse_args()

    manifest = generate(n=args.n, seed=args.seed)
    print(f"Сгенерировано резюме: {len(manifest)}")
    print(f"PDF/DOCX: {config.resumes_dir}")
    print(f"Манифест: {config.manifest_path}")


if __name__ == "__main__":
    main()
